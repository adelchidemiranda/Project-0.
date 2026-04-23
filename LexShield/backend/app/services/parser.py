"""
Document Parser Service
Supports PDF, DOCX, TXT → normalized text + structured sections
"""
import re
from pathlib import Path
from dataclasses import dataclass, field


@dataclass
class Section:
    section_id: str
    title: str
    paragraphs: list[str] = field(default_factory=list)
    char_start: int = 0
    char_end: int = 0


@dataclass
class ParsedDocument:
    raw_text: str
    normalized_text: str
    sections: list[Section]
    sections_dict: dict


def parse_document(file_path: str) -> ParsedDocument:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        raw_text = _parse_pdf(file_path)
    elif ext == ".docx":
        raw_text = _parse_docx(file_path)
    else:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            raw_text = f.read()

    normalized = _normalize_text(raw_text)
    sections = _segment_sections(normalized)
    sections_dict = {s.section_id: {"title": s.title, "paragraphs": s.paragraphs} for s in sections}

    return ParsedDocument(
        raw_text=raw_text,
        normalized_text=normalized,
        sections=sections,
        sections_dict=sections_dict,
    )


def _parse_pdf(file_path: str) -> str:
    try:
        from pdfminer.high_level import extract_text
        return extract_text(file_path)
    except Exception:
        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            return f"[PDF parsing error: {e}]"


def _parse_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        return f"[DOCX parsing error: {e}]"


def _normalize_text(text: str) -> str:
    # Remove excessive whitespace while preserving paragraph breaks
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _segment_sections(text: str) -> list[Section]:
    """Split text into logical sections by detecting headings."""
    sections = []
    # Italian/English legal heading patterns
    heading_patterns = [
        r"^(Art(?:icolo)?\.?\s*\d+[^\n]*)",
        r"^(\d+[\.\)]\s+[A-Z][^\n]{5,})",
        r"^([A-Z][A-Z\s]{4,}[A-Z])\s*$",
        r"^(PREMESSE?|OGGETTO|DEFINIZIONI?|CLAUSOLE?|ALLEGAT[IO]|CONSIDERATO|VISTA?|RITENUTO|DISPONE)\b",
        r"^(WHEREAS|DEFINITIONS?|ARTICLE|CLAUSE|SCHEDULE|WHEREAS|RECITALS)\b",
    ]
    combined = re.compile("|".join(heading_patterns), re.MULTILINE | re.IGNORECASE)

    lines = text.split("\n")
    current_section = Section(section_id="S0", title="Introduction", char_start=0)
    current_paragraphs: list[str] = []
    current_text_buf: list[str] = []
    char_pos = 0
    section_idx = 0

    for line in lines:
        line_len = len(line) + 1  # +1 for newline
        if combined.match(line.strip()) and len(line.strip()) > 4:
            # Save current section
            if current_text_buf:
                current_section.paragraphs = _split_paragraphs("\n".join(current_text_buf))
                current_section.char_end = char_pos
                sections.append(current_section)
            section_idx += 1
            current_section = Section(
                section_id=f"S{section_idx}",
                title=line.strip()[:100],
                char_start=char_pos,
            )
            current_text_buf = []
        else:
            current_text_buf.append(line)
        char_pos += line_len

    # Last section
    if current_text_buf:
        current_section.paragraphs = _split_paragraphs("\n".join(current_text_buf))
        current_section.char_end = char_pos
        sections.append(current_section)

    if not sections:
        sections = [Section(
            section_id="S0",
            title="Document",
            paragraphs=_split_paragraphs(text),
            char_start=0,
            char_end=len(text),
        )]

    return sections


def _split_paragraphs(text: str) -> list[str]:
    return [p.strip() for p in text.split("\n\n") if p.strip()]
